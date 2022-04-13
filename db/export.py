from django.http import HttpResponse
from django_xhtml2pdf.utils import generate_pdf
from docx import Document
from lxml import etree
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.styles.numbers import FORMAT_CURRENCY_USD_SIMPLE
from openpyxl.utils import get_column_letter


def render_doc(context, **kwargs):
    document = (
        Document()
    )  # https://python-docx.readthedocs.io/en/latest/#what-it-can-do
    filename = kwargs.get("filename")
    item = context["item"]
    document.add_heading(item.title, 0)
    root = etree.fromstring(item.text)
    for elem in root.iter():
        if elem.tag == "h1":
            document.add_heading(elem.text, level=1)
            document.add_paragraph("")
        if elem.tag == "h2":
            document.add_heading(elem.text, level=2)
            document.add_paragraph("")
        if elem.tag == "h3":
            document.add_heading(elem.text, level=3)
            document.add_paragraph("")
        if elem.tag == "h4":
            document.add_heading(elem.text, level=4)
            document.add_paragraph("")
        if elem.tag == "h5":
            document.add_heading(elem.text, level=5)
            document.add_paragraph("")
        if elem.tag == "h6":
            document.add_heading(elem.text, level=6)
            document.add_paragraph("")
        if elem.tag == "p":
            document.add_paragraph(elem.text)
            document.add_paragraph("")

    response = HttpResponse(content_type="docx")
    response["Content-Disposition"] = "attachment; filename=%s" % filename
    document.save(response)
    return response


def render_pdf(context, **kwargs):
    filename = kwargs.get("filename")
    template = kwargs.get("template")
    response = HttpResponse(content_type="application/pdf")
    response["Content-Disposition"] = "filename=%s" % filename
    return generate_pdf(template, context=context, file_object=response)


def render_xls(context, **kwargs):
    doc_type = kwargs.get("doc_type")
    filename = kwargs.get("filename")
    if doc_type == "Independent Government Cost Estimate":
        ################################################################################
        #                                                                              #
        #  Conversation with myself                                                    #
        #  ------------------------                                                    #
        #                                                                              #
        #  Oh snap, all this Python just to                                            #
        #  put a few values into a spreadsheet                                         #
        #  without running Excel?                                                      #
        #                                                       A thousand times yes.  #
        #                                                                              #
        #                                                                              #
        #  But don't you still have to run Excel                                       #
        #  to test your output?                                                        #
        #                                                                              #
        #                                                       Soul-crushing pain     #
        #                                                       only occurs with       #
        #                                                       prolonged use of       #
        #                                                       Excel. This house      #
        #                                                       is clean.              #
        #                                                                              #
        #  Here we go again. So you're going                                           #
        #  to write all this Python                                                    #
        #  to avoid manually entering a few cells                                      #
        #  in Excel?                                                                   #
        #                                                                              #
        #                                            My kingdom to avoid Excel!        #
        #                                                                              #
        #  This is it! Finally nearing the end.                                        #
        #  Was it worth it?                                                            #
        #                                                                              #
        #                                         Weeping silent tears of joy, yes!!   #
        #                                                                              #
        #  What have we learned?                                                       #
        #                                                                              #
        #                                                       Python4life.           #
        #                                                                              #
        #            No copies of Excel were harmed during the writing                 #
        #            of this code, but many expletives were used.                      #
        #                                                                              #
        #                                                                              #
        ################################################################################
        bold = Font(bold=True)
        border = Side(border_style="thin", color="000000")
        filename = kwargs.get("filename")
        item = context["item"]
        workbook = Workbook()
        ################################################################################
        # Sheet 1                                                                      #
        ################################################################################
        sheet1 = workbook.active
        sheet1.title = "Instructions"
        # https://stackoverflow.com/a/14450572
        sheet1.append(["Instructions".upper()])
        sheet1.column_dimensions["B"].width = 111.7
        # Merge cells
        sheet1.merge_cells("A1:Z1")
        sheet1.append(
            [
                "",
                "This template is being provided as a tool to assist the acquisition workforce in developing an IGCE for a Firm Fixed Price Product or Service.",
            ]
        )
        sheet1.append([])
        sheet1.append(
            [
                "1.",
                "The exact amount of a vendor’s quote must NOT be used as the basis for a program office’s IGCE.",
            ]
        )
        sheet1.append(
            [
                "2.",
                "The program office must conduct all necessary research to compile an accurate and complete IGCE, independent of the vendor’s pricing/cost information.",
            ]
        )
        sheet1.append(
            [
                "3.",
                "The program office should use the results of the market reaearch to substaniate the IGCE.",
            ]
        )
        sheet1.append(
            [
                "4.",
                "The program office provide a narrative to document the basis of the IGCE.",
            ]
        )
        ################################################################################
        # Sheet 2                                                                      #
        ################################################################################
        sheet2 = workbook.create_sheet(title="FFP IGCE")
        sheet2.column_dimensions["A"].width = 48
        column_index = 2
        letter_start = "B"
        entries = item.time_set.all()

        #########
        # Row 1 #
        #########
        sheet2.append(["FFP IGCE"])

        #########
        # Row 2 #
        #########
        sheet2.append(["Title:".upper(), item.subject])

        #########
        # Row 3 #
        #########
        sheet2.append(["Detailed Price Summary"])
        # Bold cells, set border
        sheet2["A1"].font = bold
        sheet2["A2"].font = bold
        sheet2["A3"].font = bold
        sheet2["A3"].border = Border(bottom=border)

        #########
        # Row 4 #
        #########
        count = 1
        row_4_col_data = []
        time_set_count = item.time_set.count()
        for i in range(time_set_count):
            row_4_col_data.append("Estimate %s" % str(count))
            row_4_col_data.append("")
            row_4_col_data.append("")
            row_4_col_data.append("")
            count += 1
        row_4_col_data.insert(0, "Contract Line Item Description")
        sheet2.append(row_4_col_data)
        # Bold cell
        sheet2["A4"].font = bold
        # Fill cells
        row_4_col_num = len(row_4_col_data) - 1
        for cell in range(row_4_col_num):
            # https://stackoverflow.com/a/50209914
            sheet2[
                get_column_letter(column_index + cell) + str(sheet2.max_row)
            ].fill = PatternFill(
                start_color="D3D3D3", end_color="D3D3D3", fill_type="solid"
            )
        # Merge cells
        merge = []
        for cell in range(row_4_col_num):
            if (column_index + cell) % 4 == 1:
                merge.append(
                    letter_start
                    + str(sheet2.max_row)
                    + ":"
                    + get_column_letter(column_index + cell)
                    + str(sheet2.max_row)
                )
                letter_start = get_column_letter(column_index + cell + 1)
        for cells in merge:
            sheet2.merge_cells(cells)
        # Bold cells and set border
        sheet2[get_column_letter(column_index) + str(sheet2.max_row)].font = bold
        for cell in range(row_4_col_num):
            sheet2[
                get_column_letter(column_index + cell) + str(sheet2.max_row)
            ].border = Border(top=border, bottom=border)
            sheet2[
                get_column_letter(column_index + cell) + str(sheet2.max_row)
            ].alignment = Alignment(horizontal="center", vertical="center")
            if (column_index + cell) % 4 == 1:
                sheet2[
                    get_column_letter(column_index + cell + 1) + str(sheet2.max_row)
                ].font = bold

        #########
        # Row 5 #
        #########
        row_5_col_data = []
        for i in range(time_set_count):
            row_5_col_data.append("Quantity")
            row_5_col_data.append("Unit")
            row_5_col_data.append("Unit Price")
            row_5_col_data.append("Total Price")
        row_5_col_data.insert(0, "")
        sheet2.append(row_5_col_data)
        # Bold and center and border
        row_5_col_num = len(row_5_col_data) - 1
        for cell in range(row_5_col_num):
            sheet2[
                get_column_letter(column_index + cell) + str(sheet2.max_row)
            ].font = bold
            sheet2[
                get_column_letter(column_index + cell) + str(sheet2.max_row)
            ].alignment = Alignment(horizontal="center", vertical="center")
            sheet2[
                get_column_letter(column_index + cell) + str(sheet2.max_row)
            ].border = Border(bottom=border, right=border)
            sheet2.column_dimensions[get_column_letter(column_index + cell)].width = 12

        #########
        # Row 6 #
        #########
        row_6_col_data = []  # Blank line
        for i in range(time_set_count):
            row_6_col_data.append("")
            row_6_col_data.append("")
            row_6_col_data.append("")
            row_6_col_data.append("")
        row_6_col_data.insert(0, "")
        sheet2.append(row_6_col_data)
        # Bold cells
        row_6_col_num = len(row_6_col_data) - 1
        sheet2[get_column_letter(column_index) + str(sheet2.max_row)].font = bold
        for cell in range(row_6_col_num):
            sheet2[
                get_column_letter(column_index + cell + 1) + str(sheet2.max_row)
            ].font = bold
            sheet2[
                get_column_letter(column_index + cell) + str(sheet2.max_row)
            ].border = Border(bottom=border, right=border)

        #########
        # Row 7 #
        #########
        row_7_col_data = []
        for entry in entries:
            row_7_col_data.append(entry.quantity)
            row_7_col_data.append(entry.unit)
            row_7_col_data.append(entry.unit_price)
            row_7_col_data.append(entry.total_price)
        row_7_col_data.insert(0, item.subject)
        sheet2.append(row_7_col_data)
        # https://openpyxl.readthedocs.io/en/stable/usage.html#using-formulae
        row_7_col_num = len(row_7_col_data) - 1
        column_total = []
        for cell in range(row_7_col_num):
            if (column_index + cell) % 4 == 1:
                column_total.append(
                    get_column_letter(column_index + cell) + str(sheet2.max_row)
                )
        # Currency + formula
        for cell in range(row_7_col_num):
            if (column_index + cell) % 4 == 1:
                sheet2[
                    get_column_letter(column_index + cell) + str(sheet2.max_row)
                ].number_format = FORMAT_CURRENCY_USD_SIMPLE
                sheet2[
                    get_column_letter(column_index + cell - 1) + str(sheet2.max_row)
                ].number_format = FORMAT_CURRENCY_USD_SIMPLE
                sheet2[
                    get_column_letter(column_index + cell) + str(sheet2.max_row)
                ] = "=%s*%s" % (
                    get_column_letter(column_index + cell - 1) + str(sheet2.max_row),
                    get_column_letter(column_index + cell - 3) + str(sheet2.max_row),
                )

        # Fill cells
        for cell in range(row_7_col_num):
            if (column_index + cell) % 4 == 1:
                sheet2[
                    get_column_letter(column_index + cell) + str(sheet2.max_row)
                ].fill = PatternFill(
                    start_color="D3D3D3", end_color="D3D3D3", fill_type="solid"
                )
            sheet2[
                get_column_letter(column_index + cell) + str(sheet2.max_row)
            ].border = Border(bottom=border, right=border)

        ############
        # Row 8-15 #
        ############
        for line in range(0, 8):  # Blank lines
            row_x_col_data = []
            for i in range(time_set_count):
                row_x_col_data.append("")
                row_x_col_data.append("")
                row_x_col_data.append("")
                row_x_col_data.append("")
            row_x_col_data.insert(0, "")
            sheet2.append(row_x_col_data)
        # Fill more cells and set border + formula + currency
        row_x_col_num = len(row_x_col_data)
        for count in range(0, 8):
            for cell in range(row_x_col_num - 1):
                if (column_index + cell) % 4 == 1:
                    sheet2[
                        get_column_letter(column_index + cell)
                        + str(sheet2.max_row - count)
                    ].fill = PatternFill(
                        start_color="D3D3D3", end_color="D3D3D3", fill_type="solid"
                    )
                    sheet2[
                        get_column_letter(column_index + cell)
                        + str(sheet2.max_row - count)
                    ] = "=%s*%s" % (
                        get_column_letter(column_index + cell - 1)
                        + str(sheet2.max_row - count),
                        get_column_letter(column_index + cell - 3)
                        + str(sheet2.max_row - count),
                    )
                    sheet2[
                        get_column_letter(column_index + cell)
                        + str(sheet2.max_row - count)
                    ].number_format = FORMAT_CURRENCY_USD_SIMPLE
                sheet2[
                    get_column_letter(column_index + cell) + str(sheet2.max_row - count)
                ].border = Border(bottom=border, right=border)

        ##########
        # Row 16 #
        ##########
        sheet2.append(["Line Item Subtotal"])
        # Fill cells and set border + formula + currency
        for cell in range(row_x_col_num - 1):
            if (column_index + cell) % 4 == 1:
                line_item_subtotal_formula = "=SUM("
                for count in range(0, 9):
                    line_item_subtotal_formula += "%s+" % (
                        get_column_letter(column_index + cell)
                        + str(sheet2.max_row - count - 1)
                    )
                line_item_subtotal_formula = line_item_subtotal_formula[:-1]
                line_item_subtotal_formula += ")"
                sheet2[
                    get_column_letter(column_index + cell) + str(sheet2.max_row)
                ].fill = PatternFill(
                    start_color="D3D3D3", end_color="D3D3D3", fill_type="solid"
                )
                sheet2[
                    get_column_letter(column_index + cell) + str(sheet2.max_row)
                ] = line_item_subtotal_formula
                sheet2[
                    get_column_letter(column_index + cell) + str(sheet2.max_row)
                ].number_format = FORMAT_CURRENCY_USD_SIMPLE
            else:
                sheet2[
                    get_column_letter(column_index + cell) + str(sheet2.max_row)
                ].fill = PatternFill(
                    start_color="00008B", end_color="00008B", fill_type="solid"
                )
            sheet2[
                get_column_letter(column_index + cell) + str(sheet2.max_row)
            ].border = Border(bottom=border, right=border)

        ##########
        # Row 17 #
        ##########
        sheet2.append(["Total Estimated Amount"])
        # Bold cell
        sheet2["A" + str(sheet2.max_row)].font = bold
        # Fill cells
        for cell in range(row_x_col_num - 1):
            if (column_index + cell) % 4 == 1:
                sheet2[
                    get_column_letter(column_index + cell) + str(sheet2.max_row)
                ].fill = PatternFill(
                    start_color="D3D3D3", end_color="D3D3D3", fill_type="solid"
                )
                sheet2[
                    get_column_letter(column_index + cell) + str(sheet2.max_row)
                ] = "=%s" % get_column_letter(column_index + cell) + str(
                    sheet2.max_row - 1
                )
                sheet2[
                    get_column_letter(column_index + cell) + str(sheet2.max_row)
                ].number_format = FORMAT_CURRENCY_USD_SIMPLE
            else:
                sheet2[
                    get_column_letter(column_index + cell) + str(sheet2.max_row)
                ].fill = PatternFill(
                    start_color="00008B", end_color="00008B", fill_type="solid"
                )
            # Border
            sheet2[
                get_column_letter(column_index + cell) + str(sheet2.max_row)
            ].border = Border(bottom=border, right=border)

        ##########
        # Row 18 #
        ##########
        sheet2.append(["Total Combined Amount".upper()])
        # Bold cell
        sheet2["A" + str(sheet2.max_row)].font = bold
        # Fill cells
        for cell in range(row_x_col_num - 1):
            sheet2[
                get_column_letter(column_index + cell) + str(sheet2.max_row)
            ].fill = PatternFill(
                start_color="00FF00", end_color="00FF00", fill_type="solid"
            )
            # Border
            sheet2[
                get_column_letter(column_index + cell) + str(sheet2.max_row)
            ].border = Border(bottom=border, right=border)
        # https://openpyxl.readthedocs.io/en/stable/usage.html#using-formulae
        sheet2["B" + str(sheet2.max_row)] = "=SUM(%s)" % "+".join(column_total)
        sheet2["B" + str(sheet2.max_row)].number_format = FORMAT_CURRENCY_USD_SIMPLE
        # Merge cells
        letter_start = "B"
        merge = []
        count = 0
        for cell in range(row_x_col_num - 1):
            if (column_index + cell) % 4 == 1:
                count += 1
                merge.append(
                    letter_start
                    + str(sheet2.max_row)
                    + ":"
                    + get_column_letter(column_index + cell)
                    + str(sheet2.max_row)
                )
                letter_start = get_column_letter(column_index + cell + 1)
        for cells in merge:
            sheet2.merge_cells(cells)

        ##########
        # Row 19 #
        ##########
        sheet2.append(["Total Average Value".upper()])
        # Bold cell
        sheet2["A" + str(sheet2.max_row)].font = bold
        # Fill cells
        for cell in range(row_x_col_num - 1):
            sheet2[
                get_column_letter(column_index + cell) + str(sheet2.max_row)
            ].fill = PatternFill(
                start_color="00FF00", end_color="00FF00", fill_type="solid"
            )
            # Border
            sheet2[
                get_column_letter(column_index + cell) + str(sheet2.max_row)
            ].border = Border(bottom=border, right=border)

        sheet2["E" + str(sheet2.max_row)] = "=SUM((%s)/%s)" % (
            "+".join(column_total),
            count,
        )
        sheet2["E" + str(sheet2.max_row)].number_format = FORMAT_CURRENCY_USD_SIMPLE

        ##########
        # Row 20 #
        ##########
        sheet2.append(["Narrative:"])
        # Bold cell
        sheet2["A" + str(sheet2.max_row)].font = bold
        ##########
        # Row 21 #
        ##########
        sheet2.append([""])
        ##########
        # Row 22 #
        ##########
        notes = item.note.all()  # Get narrative from first related note title
        if len(notes) > 0:
            sheet2.append([notes[0].title])
        else:
            sheet2.append(
                [
                    "The government estimates the cost of the Confocal Laser Scanning Microscope with the features essential to the programs needs is $26730.\n\nThe estimate is based upon the comparison the published commercial price for a Confocal Laser Scanning Microscope of similar features and functionality from three (3) major manufacturers."
                ]
            )
        ##########
        # Row 23 #
        ##########
        sheet2.append([""])
        #############
        # Row 24-31 #
        #############
        count = 1
        for entry in entries:
            sheet2.append(["Estimate %s—%s" % (str(count), entry.description)])
            sheet2["A" + str(sheet2.max_row)].font = bold
            sheet2.append([item.subject])
            sheet2.append([entry.total_price])
            sheet2["A" + str(sheet2.max_row)].number_format = FORMAT_CURRENCY_USD_SIMPLE
            count += 1
        sheet2["B2"].font = bold
        response = HttpResponse(content_type="xlsx")
        response["Content-Disposition"] = "attachment; filename=%s" % filename
        workbook.active = 1
        workbook.save(response)
        return response
    else:
        # https://openpyxl.readthedocs.io/en/stable/usage.html#write-a-workbook
        filename = kwargs.get("filename")
        workbook = Workbook()
        item = context["item"]
        sheet1 = workbook.active
        sheet1.title = "Invoice"
        for entry in item.time_set.all():
            if entry.task:
                sheet1.append(
                    [
                        entry.date,
                        entry.task.name,
                        entry.description,
                        entry.hours,
                        entry.task.rate,
                    ]
                )
            else:
                sheet1.append(
                    [
                        entry.date,
                        "Task name",
                        entry.description,
                        entry.hours,
                        "Task rate",
                    ]
                )
        response = HttpResponse(content_type="xlsx")
        response["Content-Disposition"] = "attachment; filename=%s" % filename
        workbook.save(response)
        return response
